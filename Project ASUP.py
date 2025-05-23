from tkinter import ttk
from tkinter import messagebox
from tkinter import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import load_workbook
import random
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Pt, RGBColor
import pandas as pd
import os
import logging

# Configure logging
logging.basicConfig(filename='substitution_list.log', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# This variable will be set according to selected date
day = None

# List of motivational quotes
quotes = [
    "Believe you can and you're halfway there.",
    "The only way to do great work is to love what you do.",
    "Success is not how high you have climbed, but how you make a positive difference to the world.",
    "Your time is limited, so don’t waste it living someone else’s life.",
    "The best time to plant a tree was 20 years ago. The second best time is now."
]

# Function to show a calendar and let user pick a date
def ask_for_date():
    from tkinter import Tk, Toplevel, Label, Button, CENTER
    from tkcalendar import Calendar

    root = Tk()
    root.withdraw()  # Hide main window

    selected_date = None

    def on_date_selected():
        nonlocal selected_date
        selected_date = cal.selection_get()
        sel_win.destroy()

    sel_win = Toplevel()
    sel_win.title("Select Date for Substitution")
    sel_win.geometry("600x560")
    sel_win.configure(bg="#e8f5e9")  # Soft blue-like background

    # Heading
    heading = Label(
        sel_win,
        text="Please Select the Date",
        font=("Arial Rounded MT Bold", 24, "bold"),
        fg="#145a32",  # Dark green
        bg="#e8f5e9",
        pady=18
    )
    heading.pack()

    # Calendar widget
    cal = Calendar(
        sel_win,
        selectmode="day",
        font="Arial 18",
        background="#b2dfdb",
        disabledbackground="#b2dfdb",
        bordercolor="#145a32",
        headersbackground="#81c784",
        normalbackground="#e8f5e9",
        weekendbackground="#e8f5e9",
        foreground="#145a32",
        normalforeground="#145a32",
        headersforeground="#fff",
        selectbackground="#388e3c",
        selectforeground="#fff",
        firstweekday="monday",
        font_day="Arial 18 bold",
        font_month="Arial 20 bold",
        font_year="Arial 18 bold"
    )
    cal.pack(pady=24)

    # Select button
    btn = Button(
        sel_win,
        text="Select Date",
        command=on_date_selected,
        font=("Arial", 17, "bold"),
        bg="#388e3c",
        fg="white",
        relief="raised",
        padx=36,
        pady=12,
        activebackground="#145a32",
        activeforeground="white"
    )
    btn.pack(pady=30)

    sel_win.grab_set()  # Modal window
    sel_win.wait_window()
    root.destroy()
    return selected_date

# Function to get the timetable of a teacher
def get_timetable(teacher_name):
    ab_timetable = {}
    wb = load_workbook(filename='T_TIMETABLE.xlsx')
    if teacher_name not in wb.sheetnames:
        logging.warning(f"Teacher '{teacher_name}' not found in the timetable.")
        return None
    sheet = wb[teacher_name]
    for i in range(2, 10):
        cells = sheet.cell(row=day + 2, column=i)
        ab_timetable[i-1] = cells.value if cells else '-'
    return ab_timetable

# Function to get the list of teachers and their codes
def get_teachers():
    df = pd.read_excel('TCODE-TNAME-TCLASSES.xlsx')
    teachers = df.set_index('Teacher Code')['Classes'].to_dict()
    teacher_names = df.set_index('Teacher Code')['Teacher Name'].to_dict()
    return teachers, teacher_names

# Function to group teachers based on the classes they teach
def group_teachers():
    teachers, teacher_names = get_teachers()
    groups = {(6, 7, 8): {}, (9, 10): {}, (11, 12): {}}
    wb = load_workbook(filename='T_TIMETABLE.xlsx')

    for code, classes in teachers.items():
        if teacher_names[code] not in wb.sheetnames:
            continue
        class_list = [int(c) for c in classes.split(',')]
        if any(c in class_list for c in [6, 7, 8]):
            groups[(6, 7, 8)][code] = teacher_names[code]
        if any(c in class_list for c in [9, 10]):
            groups[(9, 10)][code] = teacher_names[code]
        if any(c in class_list for c in [11, 12]):
            groups[(11, 12)][code] = teacher_names[code]

    return groups

# Function to check if a teacher has continuous free periods
def has_continuous_free_periods(teacher_timetable, min_free_periods=2):
    free_periods = [period for period, class_ in teacher_timetable.items() if class_ == ' ']
    free_periods.sort()
    for i in range(len(free_periods) - min_free_periods + 1):
        if all(free_periods[j] == free_periods[i] + j for j in range(min_free_periods)):
            return True
    return False

# Function to find a substitute teacher for an absent teacher
def find_substitute(absent_teacher, teacher_names, assigned_teachers, absent_teachers):
    groups = group_teachers()
    substitutes = {}

    assigned_periods = {teacher.strip(): [] for teacher in teacher_names.values()}

    excluded_teachers = [
        'AMIN MURMU', 'BINI P KURIAKOSE', 'JONES SOLOMON ROCHE',
        'KARUPPASAMY A', 'SOMASHEKHARAIAH D S', 'SR.THERESA JOSEPH'
    ]

    # Set of codes and names for fast lookup
    absent_teacher_codes = set(absent_teachers)
    absent_teacher_names = set(teacher_names[code].strip() for code in absent_teachers if code in teacher_names)

    if absent_teacher in teacher_names:
        absent_code = absent_teacher
        absent_name = teacher_names[absent_code].strip()
        absent_timetable = get_timetable(absent_name)

        if absent_timetable is None:
            return None

        for period, class_ in absent_timetable.items():
            if class_ is None or class_ == '-' or str(class_).strip() == '':
                substitutes[period] = 'Free Period'
                continue

            class_number = None
            if any(char.isdigit() for char in str(class_)):
                class_number = int(''.join(filter(str.isdigit, str(class_))))

            substitute_found = False

            for group, teachers_in_group in groups.items():
                if class_number in group:
                    for teacher_code in teachers_in_group:
                        teacher_name = teacher_names[teacher_code].strip()
                        # SKIP if teacher is absent, excluded, or already assigned or has class that period
                        if (teacher_code == absent_code or
                            teacher_name == absent_name or
                            teacher_name in excluded_teachers or
                            teacher_code in absent_teacher_codes or
                            teacher_name in absent_teacher_names):
                            continue

                        teacher_timetable = get_timetable(teacher_name)
                        if teacher_timetable is None:
                            continue

                        # Check if already assigned a sub for this period
                        if period in assigned_periods[teacher_name]:
                            continue

                        period_value = teacher_timetable.get(period)
                        # Must be empty, None, '-', or ' ' (space) to be free
                        if period_value in (None, '-', '', ' '):
                            if teacher_name not in assigned_teachers:
                                substitutes[period] = f"{teacher_name}\n{class_.strip()}"
                                assigned_teachers.append(teacher_name)
                                assigned_periods[teacher_name].append(period)
                                substitute_found = True
                                break
                if substitute_found:
                    break

            # Fallback: search in all groups (repeat the same filter conditions as above)
            if not substitute_found:
                for group, teachers_in_group in groups.items():
                    for teacher_code in teachers_in_group:
                        teacher_name = teacher_names[teacher_code].strip()
                        if (teacher_code == absent_code or
                            teacher_name == absent_name or
                            teacher_name in excluded_teachers or
                            teacher_code in absent_teacher_codes or
                            teacher_name in absent_teacher_names):
                            continue

                        teacher_timetable = get_timetable(teacher_name)
                        if teacher_timetable is None:
                            continue

                        if period in assigned_periods[teacher_name]:
                            continue

                        period_value = teacher_timetable.get(period)
                        if period_value in (None, '-', '', ' '):
                            if teacher_name not in assigned_teachers:
                                substitutes[period] = f"{teacher_name}\n{class_.strip()}"
                                assigned_teachers.append(teacher_name)
                                assigned_periods[teacher_name].append(period)
                                substitute_found = True
                                break
                    if substitute_found:
                        break

            if not substitute_found:
                substitutes[period] = 'No Substitute'
                logging.info(f"No substitute found for period {period} of {absent_name}")

    return substitutes

absent_teachers = []

# Function to center the window on the screen
def center_window(win):
    win.update_idletasks()
    width = win.winfo_width()
    height = win.winfo_height()
    x = (win.winfo_screenwidth() // 2) - (width // 2)
    y = (win.winfo_screenheight() // 2) - (height // 2)
    win.geometry('{}x{}+{}+{}'.format(width, height, x, y))

# Function to validate the number input
def validate_number_input(event, submit_button):
    try:
        value = int(event.widget.get())
        if value > 0:
            submit_button.config(state=NORMAL)
        else:
            submit_button.config(state=DISABLED)
    except ValueError:
        submit_button.config(state=DISABLED)

# Function to create the initial GUI
def gui(root):
    root.title("Absent Teachers Input")
    root.configure(bg='#f0f0f0')  # Light grey background
    Label(root, text='Enter the number of absent teachers:', font=('Arial', 18), bg='#f0f0f0').pack(pady=10)
    e1 = Entry(root, font=('Arial', 18))
    e1.pack(pady=10)
    submit_button = Button(root, text='Submit', command=lambda: root.after(200, get_absent_teachers, int(e1.get()), root), bg='#4CAF50', fg='white', font=('Arial', 14), state=DISABLED)
    submit_button.pack(pady=10)
    e1.bind("<KeyRelease>", lambda event: validate_number_input(event, submit_button))
    center_window(root)

def append_absent_teachers(entries):
    for e in entries:
        if e.get() != '':
            absent_teachers.append(int(e.get().split()[0]))

# Function to get the names of absent teachers
def get_absent_teachers(n, root):
    root.destroy()
    root = Tk()
    root.title("Teacher Names Input")
    root.geometry('800x600')
    root.configure(bg='#f0f0f0')  # Light grey background
    entries = []

    # Get the teacher names and codes
    _, teacher_names = get_teachers()
    teacher_options = [f"{code} - {name}" for code, name in teacher_names.items()]

    def validate_names():
        entered_names = [e.get() for e in entries if e.get() != '']
        if len(entered_names) == len(set(entered_names)):
            submit_button.config(state=NORMAL)
        else:
            submit_button.config(state=DISABLED)

    for i in range(n):
        Label(root, text=f'Enter the name of teacher {i+1}:', font=('Arial', 18), bg='#f0f0f0').pack(pady=10)
        combobox = ttk.Combobox(root, values=teacher_options, font=('Arial', 18))
        combobox.pack(pady=10)
        combobox.bind("<<ComboboxSelected>>", lambda event: validate_names())
        entries.append(combobox)

    submit_button = Button(root, text='Submit', command=lambda: (append_absent_teachers(entries), root.destroy()), bg='#4CAF50', fg='white', font=('Arial', 14), state=DISABLED)
    submit_button.pack(pady=10)
    center_window(root)

# Function to update the progress bar
def update_progress(progress, progress_label, quote_label, count, total):
    progress['value'] = (count / total) * 100
    progress_label.config(text=f"{int((count / total) * 100)}%")
    progress.update()

# Function to update the quote label every 5 seconds
def update_quote(quote_label):
    quote_label.config(text=random.choice(quotes))
    quote_label.after(5000, update_quote, quote_label)

# Function to show error messages in a window
def show_error(message):
    error_root = Tk()
    error_root.title("Error")
    error_root.geometry('400x200')
    error_root.configure(bg='#f0f0f0')
    Label(error_root, text="An error occurred:", font=('Arial', 14, 'bold'), bg='#f0f0f0', fg='red').pack(pady=10)
    Label(error_root, text=message, font=('Arial', 12), bg='#f0f0f0', wraplength=350).pack(pady=10)
    Button(error_root, text="Close", command=error_root.destroy, bg='#4CAF50', fg='white', font=('Arial', 12)).pack(pady=10)
    center_window(error_root)
    error_root.mainloop()

# Main function to run the program
def main():
    logging.info("Program started.")
    try:
        # Ask for the date first!
        chosen_date = ask_for_date()
        if not chosen_date:
            show_error("No date selected. Exiting.")
            return

        global day
        day = chosen_date.weekday()

        teachers, teacher_names = get_teachers()
        root = Tk()
        root.geometry('800x600')
        root.configure(bg='#f0f0f0')  # Light grey background
        gui(root)
        root.mainloop()

        # Create a new root for the progress bar
        progress_root = Tk()
        progress_root.title("Processing Substitutions")
        progress_root.geometry('500x200')
        progress_root.configure(bg='#f0f0f0')  # Light grey background
        Label(progress_root, text="Processing... Please wait.", font=('Arial', 14), bg='#f0f0f0').pack(pady=10)
        progress = ttk.Progressbar(progress_root, orient=HORIZONTAL, length=400, mode='determinate')
        progress.pack(pady=10)
        progress_label = Label(progress_root, text="0%", font=('Arial', 12), bg='#f0f0f0')
        progress_label.pack(pady=10)
        quote_label = Label(progress_root, text=random.choice(quotes), font=('Arial', 12, 'italic'), wraplength=400, bg='#f0f0f0')
        quote_label.pack(pady=10)
        center_window(progress_root)
        progress_root.update()

        # Start updating quotes every 5 seconds
        update_quote(quote_label)

        formatted_date = chosen_date.strftime("%A, %B %d, %Y")
        word_filename = f"substitution_list_{formatted_date}.docx"

        df_list = []  # Use a list to collect DataFrames for each absent teacher
        assigned_teachers = []

        # Loop over all absent teachers
        for count, teacher in enumerate(absent_teachers, start=1):
            try:
                substitute = find_substitute(teacher, teacher_names, assigned_teachers, absent_teachers)
                if substitute is not None:
                    if 'Period 9' in substitute:
                        del substitute['Period 9']
                    if not substitute:  # Check if substitute dictionary is empty
                        raise ValueError("No substitutes found.")
                    temp_df = pd.DataFrame.from_dict(substitute, orient='index', columns=[teacher_names[teacher]])
                    temp_df = temp_df.transpose()
                    df_list.append(temp_df)
                update_progress(progress, progress_label, quote_label, count, len(absent_teachers))
            except Exception as e:
                logging.error(f"Failed to create substitution list for {teacher_names[teacher]}: {e}")
                messagebox.showerror("Error", f"Failed to create substitution list for {teacher_names[teacher]}: {e}")

        if not df_list:
            progress_root.destroy()
            logging.warning("No substitutes were found.")
            show_error("No substitutes were found. Please contact Jones Sir")
            return

        # Concatenate all DataFrames in the list
        df = pd.concat(df_list)

        df.reset_index(inplace=True)
        df.columns = ['Absent Teacher'] + [f'Period {i}' for i in range(1, len(df.columns))]

        # Create a new Document
        doc = Document()

        # Set the orientation to landscape
        section = doc.sections[-1]
        section.orientation = WD_ORIENTATION.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        # Adjust margins
        section.left_margin = Pt(30)  # Adjust as needed
        section.right_margin = Pt(30)  # Adjust as needed
        section.top_margin = Pt(15)  # Adjust as needed
        section.bottom_margin = Pt(15)  # Adjust as needed

        # Add today's day and date, centered and bold
        p_date = doc.add_paragraph(formatted_date)
        p_date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_date = p_date.runs[0]
        run_date.font.size = Pt(14)
        run_date.font.bold = True

        # Add the table to the document with a built-in style for borders
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = 'Table Grid'  # Apply a built-in style with borders

        # Add the header row
        for j, col_name in enumerate(df.columns):
            cell = table.cell(0, j)
            cell.text = col_name
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)  # Set font size for the header
            run.font.bold = True  # Bold font for the header
            run.font.color.rgb = RGBColor(255, 0, 0)  # Set font color to red
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add the data rows
        for row_index, row in df.iterrows():
            for col_index, value in enumerate(row):
                cell = table.cell(int(row_index) + 1, col_index)
                cell.text = str(value)
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(10)  # Set font size for the data
                run.font.color.rgb = RGBColor(0, 0, 255)  # Set font color to blue
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Save the document
        doc.save(word_filename)

        # Open the output file automatically
        os.startfile(word_filename)

        # Destroy the progress window
        progress_root.destroy()

        logging.info("Program completed successfully.")
    except Exception as e:
        logging.error(f"An error occurred: {e}")
        show_error(f"An error occurred: {e}")

if __name__ == '__main__':
    main()
